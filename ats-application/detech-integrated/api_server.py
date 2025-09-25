from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
import os
import time
import logging
from datetime import datetime
import json
import hashlib
from werkzeug.utils import secure_filename
import sqlite3

# Document processing imports
import PyPDF2 as pdf
from docx import Document
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Configuration
app.config.update({
    'MAX_CONTENT_LENGTH': 10 * 1024 * 1024,  # 10MB max file size
    'UPLOAD_FOLDER': 'uploads',
    'SECRET_KEY': os.getenv('SECRET_KEY', 'detech-secret-key-2025')
})

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure Google Gemini AI
api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    logger.warning("GOOGLE_API_KEY not found! AI analysis will use fallback method.")
    genai = None
    model = None
else:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash')

# Initialize SQLite database
def init_db():
    conn = sqlite3.connect('detech.db')
    cursor = conn.cursor()
    
    # Create analyses table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS analyses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT NOT NULL,
            file_hash TEXT,
            job_description TEXT NOT NULL,
            match_percentage INTEGER,
            missing_keywords TEXT,
            strengths TEXT,
            recommendations TEXT,
            summary TEXT,
            ats_tips TEXT,
            category_scores TEXT,
            confidence_score REAL,
            processing_time_ms INTEGER,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            ai_model TEXT,
            raw_response TEXT
        )
    ''')
    
    conn.commit()
    conn.close()

# Initialize database on startup
init_db()

# File processing functions
def extract_text_from_file(file):
    """Extract text from PDF or DOCX file"""
    filename = secure_filename(file.filename.lower())
    
    try:
        if filename.endswith('.pdf'):
            return extract_pdf_text(file)
        elif filename.endswith('.docx'):
            return extract_docx_text(file)
        elif filename.endswith('.doc'):
            raise ValueError("DOC files require conversion to DOCX format")
        else:
            raise ValueError("Unsupported file format. Please use PDF or DOCX files.")
    except Exception as e:
        logger.error(f"Text extraction error for {filename}: {str(e)}")
        raise Exception(f"Could not extract text from {filename}: {str(e)}")

def extract_pdf_text(file):
    """Extract text from PDF file"""
    try:
        reader = pdf.PdfReader(file)
        
        if len(reader.pages) == 0:
            raise Exception("PDF file appears to be empty")
        
        text = ""
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                continue
        
        if not text.strip():
            raise Exception("No readable text found in PDF")
        
        return text.strip()
    except Exception as e:
        raise Exception(f"PDF processing error: {str(e)}")

def extract_docx_text(file):
    """Extract text from DOCX file"""
    try:
        doc = Document(file)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise Exception("No text found in DOCX file")
        
        return text.strip()
    except Exception as e:
        raise Exception(f"DOCX processing error: {str(e)}")

def process_with_ai(resume_text, job_description, filename):
    """Process resume with Google Gemini AI"""
    if not model:
        return create_fallback_analysis(resume_text, job_description)
    
    prompt = f"""
    You are an expert ATS (Applicant Tracking System) analyzer for the DETECH platform.
    
    Analyze this resume against the job description and provide a comprehensive assessment.
    
    JOB DESCRIPTION:
    {job_description}
    
    RESUME TEXT:
    {resume_text}
    
    FILENAME: {filename}
    
    Provide analysis in this JSON format:
    {{
        "match_percentage": 85,
        "missing_keywords": ["React", "Docker", "AWS"],
        "strengths": ["Python", "JavaScript", "Problem Solving"],
        "recommendations": [
            "Add specific React framework experience",
            "Include Docker containerization skills"
        ],
        "summary": "Strong candidate with solid technical foundation...",
        "ats_tips": [
            "Use exact keywords from job description",
            "Add quantifiable achievements"
        ],
        "category_scores": {{
            "technical_skills": 75,
            "experience_level": 80,
            "education": 85,
            "soft_skills": 70
        }},
        "confidence_score": 0.92
    }}
    
    Return ONLY the JSON object, no additional text.
    """
    
    try:
        start_time = time.time()
        response = model.generate_content(prompt)
        processing_time = int((time.time() - start_time) * 1000)
        
        if not response.text:
            raise Exception("Empty response from AI service")
        
        # Clean and parse response
        ai_text = response.text.strip()
        
        # Remove markdown formatting
        if ai_text.startswith('```json'):
            ai_text = ai_text.replace('```json', '').replace('```', '').strip()
        elif ai_text.startswith('```'):
            ai_text = ai_text.replace('```', '').strip()
        
        try:
            result = json.loads(ai_text)
        except json.JSONDecodeError:
            logger.error(f"JSON parsing error: {ai_text[:500]}...")
            result = create_fallback_analysis(resume_text, job_description)
            result['raw_ai_response'] = ai_text
        
        result['processing_time_ms'] = processing_time
        result['ai_model'] = 'gemini-1.5-flash'
        result['analyzed_at'] = datetime.now().isoformat()
        
        return result
        
    except Exception as e:
        logger.error(f"AI processing error: {str(e)}")
        result = create_fallback_analysis(resume_text, job_description)
        result['error'] = str(e)
        return result

def create_fallback_analysis(resume_text, job_description):
    """Create basic analysis when AI fails"""
    resume_lower = resume_text.lower()
    job_lower = job_description.lower()
    
    # Basic skill extraction
    common_skills = [
        'python', 'javascript', 'java', 'react', 'node.js', 'sql', 'aws', 
        'docker', 'kubernetes', 'git', 'agile', 'html', 'css'
    ]
    
    found_skills = [skill for skill in common_skills if skill in resume_lower]
    required_skills = [skill for skill in common_skills if skill in job_lower]
    missing_skills = [skill for skill in required_skills if skill not in found_skills]
    
    match_percentage = max(20, min(90, int((len(found_skills) / max(len(required_skills), 1)) * 100)))
    
    return {
        'match_percentage': match_percentage,
        'missing_keywords': missing_skills[:5],
        'strengths': found_skills[:5],
        'recommendations': [
            'Add more specific technical skills from the job description',
            'Include quantifiable achievements and metrics',
            'Use keywords from the job posting',
            'Highlight relevant project experience'
        ],
        'summary': f'Analysis completed using DETECH fallback method. Match score: {match_percentage}%',
        'ats_tips': [
            'Use exact keywords from job description',
            'Add quantifiable achievements',
            'Format resume for ATS compatibility'
        ],
        'category_scores': {
            'technical_skills': match_percentage,
            'experience_level': 60,
            'education': 70,
            'soft_skills': 50
        },
        'confidence_score': 0.3,
        'fallback_used': True
    }

def save_analysis_to_db(filename, file_hash, job_description, analysis_result):
    """Save analysis results to SQLite database"""
    try:
        conn = sqlite3.connect('detech.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            INSERT INTO analyses (
                filename, file_hash, job_description, match_percentage,
                missing_keywords, strengths, recommendations, summary,
                ats_tips, category_scores, confidence_score,
                processing_time_ms, ai_model, raw_response
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            filename, file_hash, job_description, analysis_result.get('match_percentage'),
            json.dumps(analysis_result.get('missing_keywords', [])),
            json.dumps(analysis_result.get('strengths', [])),
            json.dumps(analysis_result.get('recommendations', [])),
            analysis_result.get('summary', ''),
            json.dumps(analysis_result.get('ats_tips', [])),
            json.dumps(analysis_result.get('category_scores', {})),
            analysis_result.get('confidence_score', 0),
            analysis_result.get('processing_time_ms', 0),
            analysis_result.get('ai_model', 'fallback'),
            json.dumps(analysis_result)
        ))
        
        analysis_id = cursor.lastrowid
        conn.commit()
        conn.close()
        
        return analysis_id
    except Exception as e:
        logger.error(f"Database save error: {str(e)}")
        return None

# API Routes
@app.route('/')
def index():
    """Serve the main application page"""
    return send_from_directory('.', 'index.html')

@app.route('/api/health')
def health_check():
    """Health check endpoint"""
    try:
        # Test database connection
        conn = sqlite3.connect('detech.db')
        conn.close()
        db_status = 'connected'
    except:
        db_status = 'error'
    
    # Test AI service
    ai_status = 'available' if model else 'fallback'
    
    return jsonify({
        'status': 'healthy',
        'service': 'DETECH ATS API',
        'version': '1.0.0',
        'database': db_status,
        'ai_service': ai_status,
        'timestamp': datetime.now().isoformat()
    })

@app.route('/api/analyze', methods=['POST'])
def analyze_resume():
    """Main resume analysis endpoint"""
    try:
        # Validate request
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume']
        job_description = request.form.get('job_description', '').strip()
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not job_description or len(job_description) < 20:
            return jsonify({'error': 'Please provide a detailed job description'}), 400
        
        # Validate file type and size
        allowed_extensions = {'.pdf', '.docx', '.doc'}
        file_ext = os.path.splitext(file.filename.lower())[1]
        if file_ext not in allowed_extensions:
            return jsonify({'error': f'Invalid file type. Allowed: {", ".join(allowed_extensions)}'}), 400
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > app.config['MAX_CONTENT_LENGTH']:
            return jsonify({'error': 'File too large. Maximum size is 10MB'}), 400
        
        if file_size == 0:
            return jsonify({'error': 'File is empty'}), 400
        
        logger.info(f"Processing analysis for {file.filename}")
        
        # Extract text from resume
        resume_text = extract_text_from_file(file)
        
        # Calculate file hash
        file_content = file.read()
        file.seek(0)  # Reset file pointer
        file_hash = hashlib.md5(file_content).hexdigest()
        
        # Process with AI
        analysis_result = process_with_ai(
            resume_text=resume_text,
            job_description=job_description,
            filename=secure_filename(file.filename)
        )
        
        # Save to database
        analysis_id = save_analysis_to_db(
            secure_filename(file.filename),
            file_hash,
            job_description,
            analysis_result
        )
        
        # Add metadata to response
        analysis_result.update({
            'analysis_id': analysis_id,
            'filename': secure_filename(file.filename),
            'file_hash': file_hash,
            'service_version': 'detech-v1.0'
        })
        
        logger.info(f"Analysis completed for {file.filename} - Score: {analysis_result.get('match_percentage', 0)}%")
        
        return jsonify(analysis_result)
        
    except Exception as e:
        logger.error(f"Analysis error: {str(e)}")
        return jsonify({
            'error': f'Analysis failed: {str(e)}',
            'fallback_available': True
        }), 500

@app.route('/api/history')
def get_analysis_history():
    """Get analysis history"""
    try:
        conn = sqlite3.connect('detech.db')
        cursor = conn.cursor()
        
        cursor.execute('''
            SELECT id, filename, match_percentage, created_at
            FROM analyses
            ORDER BY created_at DESC
            LIMIT 50
        ''')
        
        analyses = []
        for row in cursor.fetchall():
            analyses.append({
                'id': row[0],
                'filename': row[1],
                'match_percentage': row[2],
                'created_at': row[3]
            })
        
        conn.close()
        
        return jsonify({
            'total': len(analyses),
            'analyses': analyses
        })
        
    except Exception as e:
        logger.error(f"History retrieval error: {str(e)}")
        return jsonify({'error': 'Could not retrieve history'}), 500

@app.route('/api/analysis/<int:analysis_id>')
def get_analysis_details(analysis_id):
    """Get detailed analysis results"""
    try:
        conn = sqlite3.connect('detech.db')
        cursor = conn.cursor()
        
        cursor.execute('SELECT * FROM analyses WHERE id = ?', (analysis_id,))
        row = cursor.fetchone()
        
        if not row:
            return jsonify({'error': 'Analysis not found'}), 404
        
        # Convert row to dict
        columns = [description[0] for description in cursor.description]
        analysis = dict(zip(columns, row))
        
        # Parse JSON fields
        for field in ['missing_keywords', 'strengths', 'recommendations', 'ats_tips', 'category_scores']:
            if analysis[field]:
                try:
                    analysis[field] = json.loads(analysis[field])
                except:
                    analysis[field] = []
        
        conn.close()
        
        return jsonify(analysis)
        
    except Exception as e:
        logger.error(f"Analysis detail error: {str(e)}")
        return jsonify({'error': 'Could not retrieve analysis details'}), 500

# Error handlers
@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 10MB.'}), 413

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(500)
def internal_error(e):
    return jsonify({'error': 'Internal server error'}), 500

if __name__ == '__main__':
    logger.info("🚀 Starting DETECH ATS API Server")
    logger.info(f"📍 Server: http://localhost:5001")
    logger.info(f"🤖 AI Service: {'Google Gemini' if model else 'Fallback Mode'}")
    logger.info(f"📁 Upload folder: {app.config['UPLOAD_FOLDER']}")
    logger.info(f"💾 Database: SQLite (detech.db)")
    
    app.run(
        host='0.0.0.0',
        port=5001,
        debug=os.getenv('FLASK_ENV') == 'development'
    )